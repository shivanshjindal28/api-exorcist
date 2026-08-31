"""
Generate submittable Word versions of the project documents.

Why a converter rather than hand-authored .docx
------------------------------------------------
The Markdown files under docs/ are the single source of truth. Maintaining a
separate hand-built Word copy guarantees the two drift apart, and the version a
reviewer reads would slowly stop matching the version the repository documents.
This script regenerates the .docx from the Markdown, so there is exactly one
place to edit.

The converter handles only the subset of Markdown these documents actually use;
it is not a general-purpose implementation and does not try to be.

Usage:
    python docs/build_docx.py                 # build every document
    python docs/build_docx.py literature      # build one

Requires: python-docx. Mermaid diagrams are embedded as images when
docs/diagrams/<name>.png exists (see render_diagrams.py); otherwise a labelled
placeholder is written so the omission is visible rather than silent.
"""

from __future__ import annotations

import re
import sys
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

DOCS = Path(__file__).resolve().parent
DIAGRAMS = DOCS / "diagrams"

BODY_FONT = "Times New Roman"
MONO_FONT = "Consolas"
INK = RGBColor(0x1C, 0x1F, 0x26)
MUTED = RGBColor(0x5A, 0x5E, 0x66)
ACCENT = RGBColor(0x8A, 0x4B, 0x12)


# ---------------------------------------------------------------------------
# Document metadata
# ---------------------------------------------------------------------------
@dataclass(frozen=True)
class DocSpec:
    key: str
    source: str
    output: str
    title: str
    subtitle: str


SPECS = [
    DocSpec(
        key="literature",
        source="literature-review.md",
        output="API_Exorcist_Literature_Review.docx",
        title="Literature Review",
        subtitle=(
            "API Exorcist — Autonomous Discovery and Safe Elimination of "
            "Zombie, Shadow and Orphaned APIs"
        ),
    ),
    DocSpec(
        key="design",
        source="design-document.md",
        output="API_Exorcist_Design_Document.docx",
        title="Design Document",
        subtitle=(
            "API Exorcist — Autonomous Discovery and Safe Elimination of "
            "Zombie, Shadow and Orphaned APIs"
        ),
    ),
]


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------
def _set_cell_background(cell, hex_colour: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")          # never "solid" — renders black
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    cell._tc.get_or_add_tcPr().append(shd)


def _add_field(paragraph, instruction: str) -> None:
    """Insert a Word field code (used for the TOC and page numbers)."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and select Update Field"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, sep, placeholder, end):
        run._r.append(el)


def _add_page_numbers(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_field(footer, "PAGE")
    for run in footer.runs:
        run.font.name = BODY_FONT
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED


# ---------------------------------------------------------------------------
# Inline formatting
# ---------------------------------------------------------------------------
_INLINE = re.compile(
    r"(\*\*.+?\*\*"        # bold
    r"|\*[^*\n]+?\*"       # italic
    r"|`[^`\n]+?`"         # code
    r"|\[[^\]]+?\]\([^)]+?\))"  # link
)


def add_inline(paragraph, text: str, *, size: int = 11, bold_all: bool = False):
    """Write text into a paragraph, honouring inline Markdown."""
    for part in _INLINE.split(text):
        if not part:
            continue
        run = paragraph.add_run()
        run.font.size = Pt(size)
        run.font.color.rgb = INK
        run.bold = bold_all

        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
            run.font.name = BODY_FONT
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = MONO_FONT
            run.font.size = Pt(size - 1)
            run.font.color.rgb = ACCENT
        elif part.startswith("*") and part.endswith("*"):
            run.text = part[1:-1]
            run.italic = True
            run.font.name = BODY_FONT
        elif part.startswith("["):
            m = re.match(r"\[([^\]]+?)\]\(([^)]+?)\)", part)
            run.text = m.group(1) if m else part
            run.font.name = BODY_FONT
            run.underline = True
            run.font.color.rgb = ACCENT
        else:
            run.text = part
            run.font.name = BODY_FONT
    return paragraph


# ---------------------------------------------------------------------------
# Block-level rendering
# ---------------------------------------------------------------------------
def add_table(doc, rows: list[list[str]]) -> None:
    if not rows:
        return
    header, *body = rows
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        add_inline(para, text, size=10, bold_all=True)
        _set_cell_background(cell, "EFEDE7")

    for row in body:
        cells = table.add_row().cells
        for i, text in enumerate(row[: len(header)]):
            cells[i].text = ""
            add_inline(cells[i].paragraphs[0], text, size=10)

    doc.add_paragraph()


def add_code_block(doc, lines: list[str], language: str = "") -> None:
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(10)
    for i, line in enumerate(lines):
        run = para.add_run(line)
        run.font.name = MONO_FONT
        run.font.size = Pt(9)
        run.font.color.rgb = INK
        if i < len(lines) - 1:
            run.add_break()

    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:color"), "A9631B")
    borders.append(left)
    para._p.get_or_add_pPr().append(borders)


#: Usable area inside the page margins, portrait and landscape.
PORTRAIT_W, PORTRAIT_H = 6.3, 8.4
LANDSCAPE_W, LANDSCAPE_H = 9.0, 5.9

#: Above this width-to-height ratio a diagram is too wide to stay legible in a
#: portrait column, and gets its own landscape page instead.
WIDE_RATIO = 1.35


def _set_orientation(doc, landscape: bool) -> None:
    """Start a new section with the given page orientation."""
    from docx.enum.section import WD_ORIENT

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    portrait_w, portrait_h = Inches(8.5), Inches(11)
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = portrait_h, portrait_w
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width, section.page_height = portrait_w, portrait_h
    section.top_margin = section.bottom_margin = Inches(0.8)
    section.left_margin = section.right_margin = Inches(1.0)
    _add_page_numbers(section)


def _fit(img_w: int, img_h: int, box_w: float, box_h: float) -> Inches:
    """Largest width that keeps the whole image inside the box."""
    scale = min(box_w / img_w, box_h / img_h)
    return Inches(img_w * scale)


def add_diagram(doc, name: str, caption: str, index: int) -> None:
    """Embed a rendered Mermaid diagram, or flag its absence visibly."""
    png = DIAGRAMS / f"{name}.png"
    if png.exists():
        from PIL import Image

        with Image.open(png) as im:
            w, h = im.size
        wide = (w / h) > WIDE_RATIO

        if wide:
            # A wide diagram squeezed into a portrait column renders its labels
            # at roughly 4pt. Give it a landscape page of its own instead.
            _set_orientation(doc, landscape=True)
            width = _fit(w, h, LANDSCAPE_W, LANDSCAPE_H)
        else:
            width = _fit(w, h, PORTRAIT_W, PORTRAIT_H)

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(str(png), width=width)
    else:
        wide = False
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"[ Diagram not rendered: {name}.png missing ]")
        run.font.name = MONO_FONT
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x9B, 0x3B, 0x2E)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(f"Figure {index}: {caption}")
    run.font.name = BODY_FONT
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = MUTED

    # Return to portrait so the prose that follows is not stranded sideways.
    if wide:
        _set_orientation(doc, landscape=False)
    else:
        doc.add_paragraph()


# ---------------------------------------------------------------------------
# Markdown parsing
# ---------------------------------------------------------------------------
def strip_status_markers(text: str) -> str:
    """Replace status emoji with words — emoji render inconsistently in Word."""
    return (
        text.replace("✅", "[implemented]")
        .replace("🔵", "[in progress]")
        .replace("⬜", "[designed]")
    )


def render_markdown(doc, md: str) -> None:
    lines = md.splitlines()
    i = 0
    figure_index = 0
    # Diagram captions, in the order the fenced blocks appear.
    captions = [
        ("component", "Component diagram — the five architectural layers"),
        ("classes", "Class diagram — core domain model"),
        ("deployment", "Deployment diagram — self-hosted, single tenant"),
        ("sequence-scan", "Sequence diagram — a discovery scan"),
        ("state-lifecycle", "State machine — endpoint lifecycle"),
        ("decision-tree", "Classification decision logic"),
        ("state-safekill", "State machine — Safe Kill Simulation"),
        ("dfd", "Data flow diagram — level 1"),
        ("url-tiers", "Two-tier URL scanner"),
    ]

    while i < len(lines):
        line = strip_status_markers(lines[i].rstrip())

        # --- fenced blocks -------------------------------------------------
        if line.startswith("```"):
            lang = line[3:].strip()
            i += 1
            block: list[str] = []
            while i < len(lines) and not lines[i].startswith("```"):
                block.append(lines[i].rstrip())
                i += 1
            i += 1
            if lang == "mermaid":
                if figure_index < len(captions):
                    name, caption = captions[figure_index]
                else:
                    name, caption = f"diagram-{figure_index}", "Diagram"
                figure_index += 1
                add_diagram(doc, name, caption, figure_index)
            else:
                add_code_block(doc, block, lang)
            continue

        # --- tables --------------------------------------------------------
        if line.startswith("|") and i + 1 < len(lines) and set(
            lines[i + 1].replace("|", "").replace(" ", "")
        ) <= {"-", ":"} and lines[i + 1].strip():
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                raw = strip_status_markers(lines[i].strip())
                cells = [c.strip() for c in raw.strip("|").split("|")]
                if set("".join(cells).replace(" ", "")) <= {"-", ":"}:
                    i += 1
                    continue
                rows.append(cells)
                i += 1
            add_table(doc, rows)
            continue

        # --- headings ------------------------------------------------------
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line.lstrip("#").strip()
            if level == 1:
                i += 1
                continue  # document title handled by the title page
            heading = doc.add_heading(level=min(level - 1, 4))
            heading.paragraph_format.space_before = Pt(14 if level == 2 else 10)
            heading.paragraph_format.space_after = Pt(6)
            for run in heading.runs:
                run.text = ""
            add_inline(heading, text, size=15 if level == 2 else 12.5, bold_all=True)
            for run in heading.runs:
                run.font.color.rgb = INK
            i += 1
            continue

        # --- horizontal rule ------------------------------------------------
        if line.strip() in ("---", "***", "___"):
            i += 1
            continue

        # --- blockquote -----------------------------------------------------
        if line.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].startswith(">"):
                quote_lines.append(strip_status_markers(lines[i].lstrip("> ").rstrip()))
                i += 1
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.4)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(8)
            add_inline(para, " ".join(quote_lines), size=11)
            for run in para.runs:
                run.italic = True
            continue

        # --- lists ----------------------------------------------------------
        bullet = re.match(r"^(\s*)[-*]\s+(.*)", line)
        number = re.match(r"^(\s*)(\d+)\.\s+(.*)", line)
        if bullet or number:
            style = "List Bullet" if bullet else "List Number"
            text = bullet.group(2) if bullet else number.group(3)
            para = doc.add_paragraph(style=style)
            para.paragraph_format.space_after = Pt(3)
            add_inline(para, text, size=11)
            i += 1
            continue

        # --- blank ----------------------------------------------------------
        if not line.strip():
            i += 1
            continue

        # --- paragraph (join continuation lines) ------------------------------
        buf = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].rstrip()
            if (
                not nxt.strip()
                or nxt.startswith(("#", "|", ">", "```", "---"))
                or re.match(r"^(\s*)([-*]|\d+\.)\s+", nxt)
            ):
                break
            buf.append(strip_status_markers(nxt))
            i += 1
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.line_spacing = 1.15
        add_inline(para, " ".join(buf), size=11)


# ---------------------------------------------------------------------------
def build(spec: DocSpec) -> Path:
    md = (DOCS / spec.source).read_text(encoding="utf-8")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(11)

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)
    _add_page_numbers(section)

    # ---- title page ----
    for _ in range(5):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(spec.title)
    run.font.name = BODY_FONT
    run.font.size = Pt(30)
    run.bold = True
    run.font.color.rgb = INK

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = s.add_run(spec.subtitle)
    run.font.name = BODY_FONT
    run.font.size = Pt(13)
    run.font.color.rgb = MUTED

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "B.Tech. Computer Science and Engineering (Cyber Security)\n"
        "Mukesh Patel School of Technology Management & Engineering\n"
        "NMIMS University\n\n"
        "2026–2027"
    )
    run.font.name = BODY_FONT
    run.font.size = Pt(11)
    run.font.color.rgb = INK

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---- table of contents ----
    h = doc.add_paragraph()
    run = h.add_run("Table of Contents")
    run.font.name = BODY_FONT
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = INK
    doc.add_paragraph()
    _add_field(doc.add_paragraph(), r'TOC \o "1-3" \h \z \u')
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---- body ----
    render_markdown(doc, md)

    out = DOCS / spec.output
    doc.save(out)
    return out


def main() -> int:
    wanted = sys.argv[1:] or [s.key for s in SPECS]
    missing_diagrams = not DIAGRAMS.exists() or not any(DIAGRAMS.glob("*.png"))

    for spec in SPECS:
        if spec.key not in wanted:
            continue
        out = build(spec)
        size_kb = round(out.stat().st_size / 1024, 1)
        print(f"  built {out.name}  ({size_kb} KB)")

    if missing_diagrams:
        print()
        print("  NOTE: no rendered diagrams found in docs/diagrams/.")
        print("  Run `python docs/render_diagrams.py` first, or the design")
        print("  document will contain placeholders where figures belong.")
    print()
    print("  Open in Word and right-click the Table of Contents ->")
    print("  Update Field to populate page numbers.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
